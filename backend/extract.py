import os
import json
from PyPDF2 import PdfReader
from docx import Document
import tempfile

"""
This script processes ZIP files containing PDF or DOCX documents along with their metadata files.
It extracts text content from the documents, combines it with metadata, and saves the results
as structured JSON files. The system handles document text extraction, cleaning, and standardization
while preserving document structure and formatting.

The workflow:
1. Extract ZIP files from a downloads folder
2. Process PDF/DOCX files and their associated metadata.txt
3. Clean and standardize the extracted text
4. Save combined data as numbered JSON files
"""


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text content from a PDF file and clean it thoroughly for RAG applications.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Cleaned extracted text content with proper sentence structure.

    Raises:
        FileNotFoundError: If the PDF file does not exist.
        Exception: For other PDF reading issues.

    Example:
        text = extract_text_from_pdf('/path/to/document.pdf')
    """
    reader = PdfReader(pdf_path)

    # Extract text from each page
    extracted_text = "\n".join(
        page.extract_text() for page in reader.pages if page.extract_text()
    )

    return clean_text(extracted_text)


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from a DOCX file and clean it for RAG applications.

    Args:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: Cleaned extracted text content with proper sentence structure.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        Exception: For other DOCX reading issues.

    Example:
        text = extract_text_from_docx('/path/to/document.docx')
    """
    doc = Document(docx_path)

    # Extract text from paragraphs and tables
    text_parts = []

    # Get text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Get text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    # Join all parts and clean
    extracted_text = "\n".join(text_parts)

    return clean_text(extracted_text)


def clean_text(text: str) -> str:
    """
    Clean and standardize extracted text for better RAG processing.

    Args:
        text (str): Raw extracted text.

    Returns:
        str: Cleaned and standardized text.

    Example:
        cleaned_text = clean_text('Original messy text')
    """
    # 1. Preserve structured data format with colon (e.g., "Date: 02.12.2024")
    preserved_lines = []
    for line in text.split("\n"):
        if (
            ":" in line and len(line.split(":")[0].strip()) < 30
        ):  # Preserve short key-value pairs
            preserved_lines.append(line.strip())
        else:
            preserved_lines.append(line.strip().replace("\n", " "))

    # 2. Join all lines with space
    text = " ".join(preserved_lines)

    # 3. Fix multiple spaces
    text = " ".join(text.split())

    # 4. Ensure proper spacing after punctuation
    for punct in [".", "!", "?"]:
        text = text.replace(f"{punct} ", f"{punct} ")
        text = text.replace(f"{punct}", f"{punct} ")

    # 5. Fix common PDF artifacts
    text = text.replace(" ,", ",")
    text = text.replace(" .", ".")
    text = text.replace(" :", ":")

    # 6. Remove paragraph breaks for RAG
    text = text.replace(".\n", ". ")

    return text.strip()


def read_metadata_file(metadata_path: str) -> dict:
    """
    Read and parse the metadata from the metadata.txt file.

    Args:
        metadata_path (str): Path to the metadata.txt file.

    Returns:
        dict: Dictionary of metadata extracted from the file.

    Raises:
        FileNotFoundError: If the metadata file does not exist.
        IOError: If there are issues reading the file.

    Example:
        metadata = read_metadata_file('/path/to/metadata.txt')
    """
    metadata = {}
    with open(metadata_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        if ":" in line:  # Parse key-value pairs separated by a colon
            key, value = line.split(":", 1)
            metadata[key.strip()] = value.strip()

    return metadata


def combine_document_and_metadata(folder_path: str, file: str) -> dict:
    """
    Combine document content (PDF or DOCX) and metadata into a single JSON-like structure.

    Args:
        folder_path (str): Path to the folder containing extracted files.
        file_path (str): Path to the file to be processed.

    Returns:
        dict: Dictionary containing combined data for the folder.

    Raises:
        ValueError: If no PDF/DOCX or metadata file is found.

    Example:
        combined_data = combine_document_and_metadata('/path/to/extracted/folder', 'filename')
    """
    # Look for PDF and DOCX files
    # pdf_files = [f for f in os.listdir(folder_path) if f.endswith(".pdf")]
    # docx_files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
    # Check extentsion of file is pdf or docx. If not raise error
    if file.lower().endswith(".pdf"):
        doc_type = "pdf"
    elif file.lower().endswith(".docx"):
        doc_type = "docx"
    else:
        raise ValueError(f"Unsupported file type for {file}")
    metadata_file = os.path.join(folder_path, "metadata.txt")

    # Check if metadata file exists
    if not os.path.exists(metadata_file):
        raise ValueError(f"No metadata.txt file found in folder: {folder_path}")

    # Handle either PDF or DOCX
    if doc_type == "pdf":
        doc_path = os.path.join(folder_path, file)
        content = extract_text_from_pdf(doc_path)
    elif doc_type == "docx":
        doc_path = os.path.join(folder_path, file)
        content = extract_text_from_docx(doc_path)

    metadata = read_metadata_file(metadata_file)
    # Also add file name and type to metadata
    metadata["file_name"] = os.path.basename(doc_path)
    metadata["file_type"] = doc_type

    # Combine the data into a single dictionary
    combined_data = {
        "file_name": os.path.basename(doc_path),
        "file_type": doc_type,
        "content": content,
        "metadata": metadata,
    }

    return combined_data


def save_combined_data(
    output_folder: str, combined_data: dict, file_index: int
) -> None:
    """
    Save the combined data (document content + metadata) as a JSON file.

    Args:
        output_folder (str): Path to the folder where JSON file will be saved.
        combined_data (dict): Combined data to be saved.
        file_index (int): Unique number for the JSON file.

    Raises:
        IOError: If there are issues creating the directory or saving the file.

    Example:
        save_combined_data('/path/to/output', combined_data_dict, 1)
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_path = os.path.join(output_folder, f"combined_data_{file_index}.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(combined_data, f, ensure_ascii=False, indent=2)
    print(f"Saved combined data to {output_path}")


def extract_data(temp_dir: tempfile.TemporaryDirectory):
    """
    Extract data from tempdir and return list of dictionaries.

    Args:
        temp_dir (tempfile.TemporaryDirectory): Temporary directory containing downloaded files.

    Returns:
        Python list of dictionaries containing extracted data.

    Raises:
        Exception: For any unexpected errors during processing.

    Example:
        Run this script to process ZIP files in the downloads folder.
    """

    folder_path = temp_dir
    try:
        # List files in the folder for debugging
        files_in_folder = os.listdir(folder_path)
        len_files = len(files_in_folder)
        print(f"\nProcessing folder {folder_path}:")
        print(f"Files found: {files_in_folder}")
        combined_data_list = []
        for it, file in enumerate(files_in_folder):
            if file == "metadata.txt":
                continue
            print(f"Processing file {it + 1}/{len_files}: {file}", end="\r")
            try:
                combined_data = combine_document_and_metadata(folder_path, file)
                # Check if content is empty or empty string
                if not combined_data["content"] or combined_data["content"] == "":
                    print(f"\nSkipping file {file} due to empty content.")
                    continue
                combined_data_list.append(combined_data)
            except Exception as e:
                print(f"\nSkipping file {file} due to error: {e}")
        return combined_data_list

    except ValueError as e:
        print(f"Error processing folder {folder_path}: {e}")
    except Exception as e:
        print(f"Unexpected error in folder {folder_path}: {str(e)}")


def main() -> None:
    # extract_data()
    pass


if __name__ == "__main__":
    main()
